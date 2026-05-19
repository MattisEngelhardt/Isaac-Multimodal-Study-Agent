import os
import tempfile
import shutil
import pytest
from unittest.mock import MagicMock, patch
import docx
from study_agent.core.processors.word_processor import WordProcessor
from study_agent.core.processors.pdf_processor import PDFProcessor

@pytest.fixture
def temp_dir():
    test_dir = tempfile.mkdtemp()
    yield test_dir
    shutil.rmtree(test_dir, ignore_errors=True)

def test_word_processor(temp_dir):
    # 1. Arrange: Create a temporary Word file
    doc_path = os.path.join(temp_dir, "test_doc.docx")
    doc = docx.Document()
    doc.add_paragraph("This is a study note paragraph.")
    
    # Add a mock table
    table = doc.add_table(rows=2, cols=2)
    table.rows[0].cells[0].text = "Header A"
    table.rows[0].cells[1].text = "Header B"
    table.rows[1].cells[0].text = "Value A"
    table.rows[1].cells[1].text = "Value B"
    doc.save(doc_path)

    # 2. Act
    wp = WordProcessor()
    text = wp.process(doc_path)

    # 3. Assert
    assert "This is a study note paragraph." in text
    assert "Header A" in text
    assert "Value B" in text
    assert "Document Table 1" in text


def test_pdf_processor_normal():
    # Arrange: Mock a normal text-based PDF page
    mock_doc = MagicMock()
    mock_page = MagicMock()
    mock_page.get_text.return_value = "Normal text content inside lecture."
    mock_doc.load_page.return_value = mock_page
    mock_doc.__len__.return_value = 1
    
    with patch('fitz.open', return_value=mock_doc), \
         patch('os.path.exists', return_value=True):
        # Act
        processor = PDFProcessor()
        text = processor.process("dummy_lecture.pdf")
        
        # Assert
        assert "Normal text content inside lecture." in text
        mock_page.get_text.assert_called_once()


def test_pdf_processor_ocr_fallback():
    # Arrange: Mock a scanned PDF page (returns empty text)
    mock_doc = MagicMock()
    mock_page = MagicMock()
    mock_page.get_text.return_value = ""  # empty trigger
    
    # Mock rendering pixmap
    mock_pix = MagicMock()
    mock_pix.tobytes.return_value = b"fake_png_data"
    mock_page.get_pixmap.return_value = mock_pix
    
    mock_doc.load_page.return_value = mock_page
    mock_doc.__len__.return_value = 1
    
    # Mock Anthropic vision response
    mock_client = MagicMock()
    mock_response = MagicMock()
    mock_text_block = MagicMock()
    mock_text_block.text = "OCR transcribed handwriting from slide"
    mock_response.content = [mock_text_block]
    mock_client.messages.create.return_value = mock_response

    with patch('fitz.open', return_value=mock_doc), \
         patch('os.path.exists', return_value=True):
        # Act
        processor = PDFProcessor(vision_client_fallback=mock_client)
        text = processor.process("scanned_notes.pdf")
        
        # Assert
        assert "OCR transcribed handwriting from slide" in text
        mock_page.get_pixmap.assert_called_once()
        mock_client.messages.create.assert_called_once()


@patch('google.generativeai.GenerativeModel')
@patch('google.generativeai.configure')
def test_pdf_processor_ocr_fallback_gemini(mock_configure, mock_model_class):
    # Arrange: Mock a scanned PDF page (returns empty text)
    mock_doc = MagicMock()
    mock_page = MagicMock()
    mock_page.get_text.return_value = ""  # empty trigger
    
    # Mock rendering pixmap
    mock_pix = MagicMock()
    mock_pix.tobytes.return_value = b"fake_png_data"
    mock_page.get_pixmap.return_value = mock_pix
    
    mock_doc.load_page.return_value = mock_page
    mock_doc.__len__.return_value = 1
    
    # Mock Gemini response
    mock_model = MagicMock()
    mock_model_class.return_value = mock_model
    mock_response = MagicMock()
    mock_response.text = "OCR transcribed handwriting from slide (Gemini)"
    mock_model.generate_content.return_value = mock_response

    with patch('fitz.open', return_value=mock_doc), \
         patch('os.path.exists', return_value=True), \
         patch('os.getenv', return_value="fake_gemini_key"):
        # Act
        processor = PDFProcessor(llm_provider="gemini", model="gemini-1.5-pro")
        text = processor.process("scanned_notes.pdf")
        
        # Assert
        assert "OCR transcribed handwriting from slide (Gemini)" in text
        mock_page.get_pixmap.assert_called_once()
        mock_configure.assert_called_once_with(api_key="fake_gemini_key")
        mock_model.generate_content.assert_called_once()
