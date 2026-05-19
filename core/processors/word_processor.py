import os
import logging
import docx

logger = logging.getLogger(__name__)

class WordProcessor:
    def __init__(self):
        pass

    def process(self, docx_path: str) -> str:
        """
        Extracts paragraphs and tables from a .docx file.
        :param docx_path: Path to the Word document.
        :return: Extracted text content.
        """
        logger.info(f"Word: Processing file {docx_path}")
        if not os.path.exists(docx_path):
            return f"[Error: Word document file {docx_path} not found]"

        try:
            doc = docx.Document(docx_path)
            full_text = []

            # 1. Read document paragraphs
            for para in doc.paragraphs:
                if para.text.strip():
                    full_text.append(para.text)

            # 2. Read document tables
            for table_index, table in enumerate(doc.tables):
                full_text.append(f"\n--- Document Table {table_index + 1} ---")
                for row in table.rows:
                    row_cells = [cell.text.strip() for cell in row.cells]
                    # Filter out empty cells, format as a pseudo Markdown table row
                    full_text.append(" | ".join(row_cells))
                full_text.append("----------------------------\n")

            return "\n".join(full_text)

        except Exception as e:
            logger.error(f"Error during Word processing: {e}")
            return f"[Error processing Word document {os.path.basename(docx_path)}: {e}]"
